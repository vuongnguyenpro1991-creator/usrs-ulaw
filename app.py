import streamlit as st
import pandas as pd
import google.generativeai as genai
import json
import time
import io
import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. CẤU HÌNH TRANG WEB
st.set_page_config(page_title="ULAW Smart Reference System", layout="wide", initial_sidebar_state="expanded")

if 'ai_data' not in st.session_state:
    st.session_state.ai_data = {
        "chuyen_nganh": ["Đang chờ AI phân tích..."],
        "tu_khoa_vn": ["Đang chờ AI phân tích..."],
        "tu_khoa_en": ["Đang chờ AI phân tích..."],
        "tu_khoa_dac_thu": ["Đang chờ AI phân tích..."],
        "co_yeu_to_nuoc_ngoai": False
    }

if 'search_clicked' not in st.session_state:
    st.session_state.search_clicked = False

if 'applied_search_mode' not in st.session_state:
    st.session_state.applied_search_mode = "Tiếng Việt & Tiếng Anh"

if 'search_results_df' not in st.session_state:
    st.session_state.search_results_df = None

if 'input_vn_widget' not in st.session_state: st.session_state.input_vn_widget = ""
if 'input_en_widget' not in st.session_state: st.session_state.input_en_widget = ""

# Các hàm xử lý giao diện (GIỮ NGUYÊN CODE CỦA BẠN)
def on_add_vn():
    val = st.session_state.input_vn_widget
    if val:
        new_kws = [k.strip() for k in val.split(',') if k.strip()]
        if st.session_state.ai_data["tu_khoa_vn"] == ["Đang chờ AI phân tích..."]: st.session_state.ai_data["tu_khoa_vn"] = []
        for kw in new_kws:
            if kw not in st.session_state.ai_data["tu_khoa_vn"]: st.session_state.ai_data["tu_khoa_vn"].append(kw)
        st.session_state.input_vn_widget = ""

def on_add_en():
    val = st.session_state.input_en_widget
    if val:
        new_kws = [k.strip() for k in val.split(',') if k.strip()]
        if st.session_state.ai_data["tu_khoa_en"] == ["Đang chờ AI phân tích..."]: st.session_state.ai_data["tu_khoa_en"] = []
        for kw in new_kws:
            if kw not in st.session_state.ai_data["tu_khoa_en"]: st.session_state.ai_data["tu_khoa_en"].append(kw)
        st.session_state.input_en_widget = ""

def call_gemini(topic, mode):
    try:
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
        model = genai.GenerativeModel('gemini-1.5-flash') # Đã sửa thành model ổn định trên cloud
        mode_instruction = ""
        if mode == "Tiếng Việt":
            mode_instruction = """
5. PHẠM VI NGÔN NGỮ: Người dùng CHỈ yêu cầu tra cứu bằng Tiếng Việt.
- "tu_khoa_en" BẮT BUỘC là mảng rỗng [].
- "tu_khoa_dac_thu" BẮT BUỘC là mảng rỗng [].
- "co_yeu_to_nuoc_ngoai" BẮT BUỘC là false.
"""
        else:
            mode_instruction = """
5. PHẠM VI NGÔN NGỮ: Dịch các từ khóa cốt lõi sang Tiếng Anh để đưa vào "tu_khoa_en". Nếu có yếu tố đặc thù quốc tế, hãy cập nhật "tu_khoa_dac_thu".
"""
        prompt = f"""Bạn là một Chuyên gia Thư viện học tại Đại học Luật TP.HCM. 
Nhiệm vụ: Phân tích đề tài sau để tạo bộ từ khóa tra cứu: "{topic}".
KỶ LUẬT NGHIỆP VỤ:
1. Từ khóa ĐẦU TIÊN phải là TỪ KHÓA CHÍNH mang tính cốt lõi của đề tài.
2. TUYỆT ĐỐI KHÔNG LIỆT KÊ TỪ ĐỒNG NGHĨA. 
3. Các từ khóa tiếp theo PHẢI LÀ TỪ KHÓA PHỤ.
4. CẤM tạo ra các từ khóa quá dài ghép nhiều vế.{mode_instruction}
Trả về ĐÚNG JSON: {{"chuyen_nganh": ["..."], "tu_khoa_vn": ["..."], "tu_khoa_en": ["..."], "tu_khoa_dac_thu": ["..."], "co_yeu_to_nuoc_ngoai": true/false}} không giải thích."""
        
        response = model.generate_content(prompt)
        text_result = response.text.replace("```json", "").replace("```", "").strip()
        st.session_state.ai_data = json.loads(text_result)
        return True
    except Exception as e:
        st.error(f"Lỗi hệ thống AI: Vui lòng kiểm tra lại API Key.")
        return False

# BỔ SUNG: THUẬT TOÁN TÌM KIẾM CÓ TÍCH HỢP LỌC TRÙNG & LOCAL-FIRST
def generate_mock_results(topic, selected_sources):
    """Giả lập Data có cấu trúc chuẩn để đưa vào file Word"""
    raw_data = [
        {"Tên tài liệu": f"Pháp luật về {topic[:30]}", "Tác giả": "Nguyễn Văn A", "Năm": 2023, "Nguồn": "Thư viện ĐH Luật TP.HCM", "Loại hình": "Luận văn", "Nơi xuất bản": "TP.HCM", "Nhà xuất bản": "", "Người hướng dẫn": "PGS.TS Lê B", "Số trang": "85", "Tạp chí": "", "Số kỳ": "", "Link": ""},
        {"Tên tài liệu": f"Pháp luật về {topic[:30]}", "Tác giả": "Nguyễn Văn A", "Năm": 2023, "Nguồn": "Thư viện ĐH Luật HN", "Loại hình": "Luận văn", "Nơi xuất bản": "TP.HCM", "Nhà xuất bản": "", "Người hướng dẫn": "Lê B", "Số trang": "85", "Tạp chí": "", "Số kỳ": "", "Link": ""}, # Bản ghi trùng lặp
        {"Tên tài liệu": f"Thực trạng {topic[:20]}", "Tác giả": "Trần Thị C", "Năm": 2024, "Nguồn": "Thư viện UEL", "Loại hình": "Bài báo", "Nơi xuất bản": "", "Nhà xuất bản": "", "Người hướng dẫn": "", "Số trang": "12-18", "Tạp chí": "Khoa học Pháp lý", "Số kỳ": "02", "Link": ""},
        {"Tên tài liệu": f"Nghị định quy định về {topic[:15]}", "Tác giả": "Chính phủ", "Năm": 2020, "Nguồn": "Thư Viện Pháp Luật", "Loại hình": "Văn bản pháp quy", "Nơi xuất bản": "", "Nhà xuất bản": "", "Người hướng dẫn": "", "Số trang": "", "Tạp chí": "", "Số kỳ": "", "Link": "https://tvpl.vn/.."},
        {"Tên tài liệu": "Comparative Analysis in Law", "Tác giả": "John Doe", "Năm": 2022, "Nguồn": "HeinOnline", "Loại hình": "Bài báo", "Nơi xuất bản": "", "Nhà xuất bản": "", "Người hướng dẫn": "", "Số trang": "45-60", "Tạp chí": "Harvard Law Review", "Số kỳ": "1", "Link": ""},
    ]
    df = pd.DataFrame(raw_data)
    
    # Lọc theo nguồn được user chọn
    if selected_sources:
        # Ánh xạ tên để khớp với danh sách nguồn
        source_mapping = {"Thư viện Trường Đại học Luật TP.HCM": "Thư viện ĐH Luật TP.HCM", "Thư viện Trường Đại học Luật Hà Nội": "Thư viện ĐH Luật HN", "Thư viện Trường Đại học Kinh tế - Luật": "Thư viện UEL"}
        selected_mapped = [source_mapping.get(s, s) for s in selected_sources]
        df = df[df['Nguồn'].isin(selected_mapped)]
    
    if df.empty: return df

    # Local-first Routing: Ưu tiên ULAW
    df['Priority'] = df['Nguồn'].apply(lambda x: 1 if "Luật TP.HCM" in x else 2)
    df = df.sort_values('Priority').drop_duplicates(subset=['Tên tài liệu', 'Tác giả'], keep='first')
    
    # Thêm cột Checkbox để user kiểm duyệt
    df.insert(0, 'Chọn', True)
    return df.drop(columns=['Priority'])

# BỔ SUNG: XUẤT FILE DOCX CHUẨN ULAW THEO CẤU TRÚC 4 PHẦN VÀ BẢNG THỐNG KÊ
def format_citation(row):
    """Định dạng chuỗi tĩnh theo chuẩn ISBD của ULAW"""
    t = row.get('Loại hình', '')
    title = str(row.get('Tên tài liệu', '')).strip()
    author = str(row.get('Tác giả', '')).strip()
    year = str(row.get('Năm', '')).strip()
    
    def get_val(col): return str(row.get(col, "")).strip() if pd.notna(row.get(col, "")) else ""

    pages = f", {get_val('Số trang')} tr." if get_val('Số trang') else ""
    guide = f"; Người hướng dẫn: {get_val('Người hướng dẫn')}" if get_val('Người hướng dẫn') else ""
    place = f", {get_val('Nơi xuất bản')}" if get_val('Nơi xuất bản') else ""
    publisher = f", {get_val('Nhà xuất bản')}" if get_val('Nhà xuất bản') else ""
    journal = f", {get_val('Tạp chí')}" if get_val('Tạp chí') else ""
    issue = f", Số {get_val('Số kỳ')}" if get_val('Số kỳ') else ""
    link = f" {get_val('Link')} (Truy cập ngày {datetime.datetime.now().strftime('%d/%m/%Y')})" if get_val('Link') else ""

    if t in ["Luận án", "Luận văn", "Khóa luận tốt nghiệp", "Khóa luận"]:
        return f"{title} : {t}, {author}{guide}{place}, {year}{pages}"
    elif t in ["Bài báo", "Tạp chí"]:
        return f"{title}, {author}{journal}{publisher}, {year}{issue}{pages}"
    elif t == "Sách":
        return f"{title}, {author}{place}{publisher}, {year}{pages}"
    elif t == "Văn bản pháp quy":
        return f"{title}, Cơ quan ban hành: {author}, Năm ban hành: {year}{link}"
    else:
        return f"{title}, {author}, {year}{pages}"

def generate_word_report(df, topic, author, tk_vn, tk_en):
    doc = Document()
    for section in doc.sections:
        section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8)

    # Header
    table = doc.add_table(rows=1, cols=2)
    p_left = table.cell(0, 0).paragraphs[0]
    p_left.add_run("TRƯỜNG ĐẠI HỌC LUẬT TP.HCM\nTHƯ VIỆN\n-------").bold = True
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_right = table.cell(0, 1).paragraphs[0]
    p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n--------------------").bold = True
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("KẾT QUẢ TRA CỨU THÔNG TIN")
    title_run.bold = True
    title_run.font.size = Pt(15)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Đề tài: “{topic}”").runs[0].bold = True
    doc.add_paragraph(f"Từ khóa tiếng Việt: “{', '.join(tk_vn)}”")
    doc.add_paragraph(f"Từ khóa tiếng Anh: “{', '.join(tk_en)}”")
    doc.add_paragraph(f"Kính gửi tác giả hoặc nhóm tác giả: {author}").runs[0].bold = True
    doc.add_paragraph("Sau khi tiến hành các quy trình tra cứu khảo sát thông tin đề tài nghiên cứu, Thư viện Trường ĐH Luật TP. Hồ Chí Minh kính gửi đến anh, chị danh mục tài liệu có liên quan đến đề tài sau đây:")
    
    # Sắp xếp và in Phần A
    df_sorted = df.sort_values(by=['Năm', 'Tên tài liệu'], ascending=[False, True])
    doc.add_heading("PHẦN A: CÁC WEBSITE PHẢI SỬ DỤNG TRONG TRA CỨU ĐỐI VỚI TẤT CẢ ĐỀ TÀI", level=1)
    
    academic_sources = df_sorted[~df_sorted['Loại hình'].isin(['Văn bản pháp quy', 'Website chuyên ngành'])]['Nguồn'].unique()
    type_order = ["Luận án", "Luận văn", "Khóa luận tốt nghiệp", "Đề tài nghiên cứu khoa học", "Sách", "Bài báo"]
    
    for idx, source in enumerate(academic_sources, 1):
        doc.add_paragraph(f"{idx}. {source}:").runs[0].bold = True
        source_df = df_sorted[df_sorted['Nguồn'] == source]
        for t in type_order:
            type_df = source_df[source_df['Loại hình'] == t]
            p = doc.add_paragraph(f"    • {t}: ")
            if type_df.empty:
                p.add_run("Không có tài liệu").italic = True
            else:
                for _, row in type_df.iterrows():
                    doc.add_paragraph(format_citation(row), style='List Bullet 2')
                    
    # Phần B
    doc.add_heading("PHẦN B: VĂN BẢN PHÁP QUY LIÊN QUAN ĐẾN ĐỀ TÀI", level=1)
    vb_df = df_sorted[df_sorted['Loại hình'] == 'Văn bản pháp quy']
    if vb_df.empty: doc.add_paragraph("    Không có tài liệu").italic = True
    else:
        for _, row in vb_df.iterrows(): doc.add_paragraph(format_citation(row), style='List Bullet')

    # Phần D: Bảng thống kê
    doc.add_heading("PHẦN D: THỐNG KÊ KẾT QUẢ KHẢO SÁT", level=1)
    stat_table = doc.add_table(rows=1, cols=3)
    stat_table.style = 'Table Grid'
    hdr_cells = stat_table.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text = 'STT', 'Loại hình tài liệu', 'Số lượng'
    
    total_docs = 0
    all_types = type_order + ['Văn bản pháp quy', 'Website chuyên ngành']
    for i, t in enumerate(all_types, 1):
        count = len(df_sorted[df_sorted['Loại hình'] == t])
        total_docs += count
        row_cells = stat_table.add_row().cells
        row_cells[0].text, row_cells[1].text, row_cells[2].text = str(i), t, f"{count:02d}"
        
    row_total = stat_table.add_row().cells
    row_total[1].text = "TỔNG"
    row_total[1].paragraphs[0].runs[0].bold = True
    row_total[2].text = f"{total_docs:02d}"
    
    # Chữ ký
    doc.add_paragraph("\n")
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.add_run(f"TP. Hồ Chí Minh, ngày {datetime.datetime.now().day:02d} tháng {datetime.datetime.now().month:02d} năm {datetime.datetime.now().year}\n").italic = True
    footer_p.add_run("GIÁM ĐỐC THƯ VIỆN\n\n\n\n").bold = True
    footer_p.add_run("Ngô Kim Hoàng Nguyên").bold = True

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# 2. TÙY CHỈNH CSS CHUYÊN SÂU (GIỮ NGUYÊN)
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    html, body, [class*="css"] {font-family: 'Inter', sans-serif;}
    .stApp {background-color: #f4f6f9;}
    header {visibility: hidden;}
    [data-testid="stSidebar"] {background-color: #2b3b7c; color: white;}
    [data-testid="stSidebar"] * {color: white !important;}
    .hero-banner {
        background-image: linear-gradient(rgba(26, 35, 126, 0.85), rgba(26, 35, 126, 0.85)), url('https://images.unsplash.com/photo-1507842217343-583bb7270b66?q=80&w=2000&auto=format&fit=crop');
        background-size: cover; background-position: center;
        padding: 40px 30px; border-radius: 10px; margin-top: -50px; margin-bottom: 20px;
        text-align: center; box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    .hero-title-small {color: #90caf9; font-weight: 700; font-size: 16px; letter-spacing: 1.5px; text-transform: uppercase;}
    .hero-title-large {color: #ffffff; font-weight: 800; font-size: 36px; margin: 15px 0; text-shadow: 1px 1px 3px rgba(0,0,0,0.3);}
    .card-header {
        font-weight: 700; color: #1a237e; font-size: 16px; margin-bottom: 20px; 
        display: flex; align-items: center; background-color: #e8eaf6; padding: 12px 15px; border-radius: 6px; border-left: 5px solid #2962ff;
    }
    textarea {background-color: #0f172a !important; color: #bae6fd !important; font-family: 'Courier New', Courier, monospace !important; font-size: 13.5px !important; border-left: 4px solid #38bdf8 !important; line-height: 1.5 !important;}
    .step-container {display: flex; justify-content: space-between; background: white; padding: 15px 30px; border-radius: 10px; margin-bottom: 20px; border: 1px solid #e0e0e0;}
    .step-item {display: flex; align-items: center; font-weight: 600; color: #5f6368; font-size: 14px;}
    .step-circle {background-color: #e8eaf6; color: #3f51b5; width: 30px; height: 30px; border-radius: 50%; display: flex; justify-content: center; align-items: center; margin-right: 10px; font-weight: bold;}
    .step-active .step-circle {background-color: #3f51b5; color: white;}
    .step-active {color: #3f51b5;}
    .btn-blue button {background-color: #2962ff !important; color: white !important; width: 100%; border-radius: 8px !important; font-weight: 700 !important; border: none; padding: 10px !important;}
    .btn-green button {background-color: #00897b !important; color: white !important; width: 100%; border-radius: 8px !important; font-weight: 700 !important; border: none; padding: 10px !important;}
    .btn-orange button {background-color: #f57c00 !important; color: white !important; width: 100%; border-radius: 8px !important; font-weight: 700 !important; border: none; padding: 10px !important;}
    .btn-outline button {background-color: white !important; color: #3f51b5 !important; border: 1.5px solid #3f51b5 !important; width: 100%; border-radius: 8px !important; font-weight: 600 !important;}
    div[data-testid="stVerticalBlock"] > div[style*="flex-direction: column;"] > div[data-testid="stVerticalBlock"] {background-color: white; padding: 20px; border-radius: 10px; border: 1px solid #e0e0e0; box-shadow: 0 2px 8px rgba(0,0,0,0.03);}
    .scrollable-source {max-height: 380px; overflow-y: auto; padding-right: 10px;}
    
    /* STYLE MỚI CHO CÁC TIÊU ĐỀ PHÂN NHÓM CỘT 3 */
    .source-category-title {
        font-weight: 700; color: #1e293b; 
        margin-top: 15px; margin-bottom: 5px; 
        border-bottom: 1px solid #cbd5e1; padding-bottom: 3px; 
        font-size: 13px; text-transform: uppercase;
    }
</style>
""", unsafe_allow_html=True)

# 3. SIDEBAR (GIỮ NGUYÊN)
with st.sidebar:
    col_logo1, col_logo2, col_logo3 = st.columns([1,3,1])
    with col_logo2:
        try: st.image("logo_ulaw.png", use_column_width=True)
        except Exception: pass
    st.markdown("<h3 style='text-align: center;'>THƯ VIỆN ULAW</h3><hr>🏠 Trang chủ<br>📝 Tạo yêu cầu<br>⏱️ Lịch sử<br>📁 Kho báo cáo<hr>📞 (028) 3940 0989", unsafe_allow_html=True)

# HERO BANNER
st.markdown("<div class='hero-banner'><div class='hero-title-small'>ULAW SMART REFERENCE SYSTEM</div><div class='hero-title-large'>TRA CỨU THÔNG TIN THEO YÊU CẦU</div></div>", unsafe_allow_html=True)
st.markdown("<div class='step-container'><div class='step-item step-active'><div class='step-circle'>1</div> Khởi tạo yêu cầu</div><div class='step-item step-active'><div class='step-circle'>2</div> AI đề xuất từ khóa</div><div class='step-item step-active'><div class='step-circle'>3</div> Chọn nguồn</div><div class='step-item step-active'><div class='step-circle' style='background-color:#f57c00;color:white;'>4</div> Xuất báo cáo</div></div>", unsafe_allow_html=True)

# 4. BỐ CỤC 4 CỘT
col1, col2, col3, col4 = st.columns(4)

with col1:
    with st.container(border=True):
        st.markdown("<div class='card-header'>📋 1. KHỞI TẠO YÊU CẦU</div>", unsafe_allow_html=True)
        topic = st.text_input("Tên đề tài nghiên cứu *", placeholder="Nhập tên đề tài...")
        author = st.text_input("Tên người/nhóm yêu cầu *", placeholder="Nhập tên người yêu cầu...")
        st.markdown("<hr style='margin:15px 0;'>", unsafe_allow_html=True)
        
        time_mode = st.radio("📅 Khoảng thời gian xuất bản:", ["Tất cả các năm", "Giới hạn thời gian"], horizontal=True)
        if time_mode == "Giới hạn thời gian":
            c1, c2 = st.columns(2)
            with c1: year_start = st.number_input("Từ năm", value=2020, step=1)
            with c2: year_end = st.number_input("Đến năm", value=2026, step=1)
        else: year_start, year_end = None, None
            
        st.markdown("<br>", unsafe_allow_html=True)
        search_mode = st.radio("🌐 Ngôn ngữ tra cứu:", ["Tiếng Việt", "Tiếng Việt & Tiếng Anh"], horizontal=True)
        
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("<div class='btn-blue'>", unsafe_allow_html=True)
        if st.button("🚀 Phân tích bằng AI", use_container_width=True):
            if topic.strip() == "": st.error("Vui lòng nhập Tên đề tài!")
            else:
                st.session_state.search_clicked = False 
                st.session_state.search_results_df = None # Reset kết quả cũ
                st.session_state.applied_search_mode = search_mode 
                with st.spinner("🤖 AI đang phân tích dữ liệu..."):
                    if call_gemini(topic, search_mode): st.success("✅ Phân tích thành công!")
        st.markdown("</div>", unsafe_allow_html=True)

with col2:
    with st.container(border=True):
        st.markdown("<div class='card-header'>💡 2. AI ĐỀ XUẤT TỪ KHÓA</div>", unsafe_allow_html=True)
        ai = st.session_state.ai_data
        st.multiselect("Chuyên ngành luật", ai.get("chuyen_nganh", []), default=ai.get("chuyen_nganh", []))
        
        sel_vn = st.multiselect("Từ khóa tiếng Việt", ai.get("tu_khoa_vn", []), default=ai.get("tu_khoa_vn", []))
        st.text_input("➕ Thêm từ khác, sau đó nhấn Enter:", key="input_vn_widget", on_change=on_add_vn)

        if sel_vn and sel_vn[0] != "Đang chờ AI phân tích...":
            vn_text = "🔸 TÌM ĐƠN LẺ:\n"
            for kw in sel_vn: vn_text += f'"{kw}"\n'
            if len(sel_vn) > 1:
                vn_text += "\n🔸 TÌM KẾT HỢP:\n"
                core_kw = sel_vn[0]
                for sec_kw in sel_vn[1:]: vn_text += f'"{core_kw}" AND "{sec_kw}"\n'
            st.caption("✍ Lệnh tra cứu Tiếng Việt (Có thể chỉnh sửa):")
            st.text_area("Lệnh TV", value=vn_text.strip(), height=160, key="ta_vn", label_visibility="collapsed")
        
        if st.session_state.applied_search_mode == "Tiếng Việt & Tiếng Anh":
            st.markdown("<hr style='margin:15px 0; border-top: 1px dashed #cbd5e1;'>", unsafe_allow_html=True)
            sel_en = st.multiselect("Từ khóa tiếng Anh", ai.get("tu_khoa_en", []), default=ai.get("tu_khoa_en", []))
            st.text_input("➕ Thêm từ khác (Tiếng Anh), nhấn Enter:", key="input_en_widget", on_change=on_add_en)

            if sel_en and sel_en[0] != "Đang chờ AI phân tích...":
                en_text = "🔸 TÌM ĐƠN LẺ:\n"
                for kw in sel_en: en_text += f'"{kw}"\n'
                if len(sel_en) > 1:
                    en_text += "\n🔸 TÌM KẾT HỢP:\n"
                    core_kw_en = sel_en[0]
                    for sec_kw_en in sel_en[1:]: en_text += f'"{core_kw_en}" AND "{sec_kw_en}"\n'
                st.caption("✍ Lệnh tra cứu Tiếng Anh (Có thể chỉnh sửa):")
                st.text_area("Lệnh EN", value=en_text.strip(), height=160, key="ta_en", label_visibility="collapsed")
            
            sel_dac_thu = []
            if ai.get("co_yeu_to_nuoc_ngoai", False) and ai.get("tu_khoa_dac_thu", []) and ai["tu_khoa_dac_thu"][0] != "Đang chờ AI phân tích...":
                sel_dac_thu = st.multiselect("Từ khóa đặc thù (Quốc tế)", ai.get("tu_khoa_dac_thu", []), default=ai.get("tu_khoa_dac_thu", []))
                if sel_dac_thu:
                    dt_text = "🔸 TÌM ĐƠN LẺ ĐẶC THÙ:\n"
                    for kw in sel_dac_thu: dt_text += f'"{kw}"\n'
                    if sel_vn and sel_vn[0] != "Đang chờ AI phân tích...":
                        dt_text += "\n🔸 ĐỐI CHIẾU BỐI CẢNH QUỐC TẾ:\n"
                        dt_text += f'"{sel_vn[0]}" AND "{sel_dac_thu[0]}"'
                    st.caption("✍ Lệnh tra cứu Đặc thù (Có thể chỉnh sửa):")
                    st.text_area("Lệnh Đặc thù", value=dt_text.strip(), height=130, key="ta_dt", label_visibility="collapsed")

with col3:
    with st.container(border=True):
        st.markdown("<div class='card-header'>📚 3. CHỌN NGUỒN TRA CỨU</div>", unsafe_allow_html=True)
        st.markdown("<div class='scrollable-source'>", unsafe_allow_html=True)
        
        sources_selected = []
        
        st.markdown("<div class='source-category-title'>Thư viện đại học</div>", unsafe_allow_html=True)
        if st.checkbox("Thư viện Trường Đại học Luật TP.HCM", value=True): sources_selected.append("Thư viện Trường Đại học Luật TP.HCM")
        if st.checkbox("Thư viện Trường Đại học Kinh tế - Luật", value=True): sources_selected.append("Thư viện Trường Đại học Kinh tế - Luật")
        if st.checkbox("Thư viện Trường Đại học Luật Hà Nội", value=True): sources_selected.append("Thư viện Trường Đại học Luật Hà Nội")
        if st.checkbox("Thư viện Trường Đại học Luật - Đại học Huế"): sources_selected.append("Thư viện ĐH Luật Huế")
        if st.checkbox("Thư viện Đại học Quốc gia Hà Nội"): sources_selected.append("Thư viện ĐHQGHN")
        if st.checkbox("Thư viện Trường Đại học Ngoại thương"): sources_selected.append("Thư viện ĐH Ngoại thương")
        if st.checkbox("Thư viện Đại học Cần Thơ"): sources_selected.append("Thư viện ĐH Cần Thơ")

        st.markdown("<div class='source-category-title'>Thư viện quốc gia và công cộng</div>", unsafe_allow_html=True)
        if st.checkbox("Thư viện Quốc gia Việt Nam", value=True): sources_selected.append("Thư viện Quốc gia Việt Nam")
        if st.checkbox("Thư viện Khoa học Tổng hợp TP.HCM"): sources_selected.append("Thư viện KHTH TP.HCM")

        st.markdown("<div class='source-category-title'>Cơ sở dữ liệu quốc tế</div>", unsafe_allow_html=True)
        if st.checkbox("HeinOnline", value=True): sources_selected.append("HeinOnline")
        if st.checkbox("Westlaw", value=True): sources_selected.append("Westlaw")
        if st.checkbox("Oxford Academic"): sources_selected.append("Oxford Academic")

        st.markdown("<div class='source-category-title'>Nguồn văn bản pháp luật</div>", unsafe_allow_html=True)
        if st.checkbox("Luật Việt Nam", value=True): sources_selected.append("Luật Việt Nam")
        if st.checkbox("Thư Viện Pháp Luật", value=True): sources_selected.append("Thư Viện Pháp Luật")

        st.markdown("<div class='source-category-title'>Các Website, CSDL chuyên ngành khác</div>", unsafe_allow_html=True)
        if st.checkbox("Google Scholar", value=True): sources_selected.append("Google Scholar")
        
        st.markdown("</div><br>", unsafe_allow_html=True)
        st.markdown("<div class='btn-green'>", unsafe_allow_html=True)
        if st.button("🔍 Bắt đầu Tra cứu", use_container_width=True):
            if not topic:
                st.error("Chưa có tên đề tài!")
            elif not sources_selected:
                st.warning("Vui lòng chọn ít nhất 1 nguồn tra cứu!")
            else:
                st.session_state.search_clicked = True
                st.session_state.search_results_df = None # Xóa dữ liệu cũ để quét lại
        st.markdown("</div>", unsafe_allow_html=True)

with col4:
    with st.container(border=True):
        st.markdown("<div class='card-header'>🏠 4. KẾT QUẢ & XUẤT BÁO CÁO</div>", unsafe_allow_html=True)
        
        if not st.session_state.search_clicked:
            st.markdown("<small style='color:#5f6368;'>Tiến trình tra cứu: <b>0%</b></small>", unsafe_allow_html=True)
            st.progress(0)
            st.info("Bấm 'Bắt đầu Tra cứu' ở Cột 3 để hệ thống tự động tổng hợp dữ liệu.")
        else:
            # QUY TRÌNH TRA CỨU ĐÃ CẬP NHẬT KIỂM DUYỆT (DATA EDITOR)
            if st.session_state.search_results_df is None:
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                # Thanh tiến trình quét qua các nguồn đã chọn
                for i, src in enumerate(sources_selected):
                    time.sleep(0.3) # Giả lập độ trễ mạng
                    progress = int(((i + 1) / len(sources_selected)) * 100)
                    progress_bar.progress(progress)
                    status_text.markdown(f"<small style='color:#5f6368;'>Đang truy vấn: <b>{src}</b>... ({progress}%)</small>", unsafe_allow_html=True)
                
                status_text.markdown("<small style='color:#f57c00;'>Đang tổng hợp và loại bỏ trùng lặp...</small>", unsafe_allow_html=True)
                st.session_state.search_results_df = generate_mock_results(topic, sources_selected)
                status_text.markdown("<small style='color:green;'><b>✓ Đã hoàn thành quét dữ liệu chéo nguồn!</b></small>", unsafe_allow_html=True)
            else:
                st.progress(100)
                st.markdown("<small style='color:green;'><b>✓ Đã hoàn thành quét dữ liệu chéo nguồn!</b></small>", unsafe_allow_html=True)
            
            df_result = st.session_state.search_results_df
            
            if df_result is not None and not df_result.empty:
                st.markdown(f"<b>Tìm thấy {len(df_result)} tài liệu. Vui lòng kiểm duyệt:</b>", unsafe_allow_html=True)
                
                # DATA EDITOR CHO PHÉP THỦ THƯ TICK CHỌN/LOẠI BỎ TÀI LIỆU
                edited_df = st.data_editor(
                    df_result,
                    column_config={"Chọn": st.column_config.CheckboxColumn("Giữ lại", default=True)},
                    disabled=["Tên tài liệu", "Tác giả", "Loại hình", "Nguồn", "Năm"],
                    hide_index=True,
                    use_container_width=True,
                    height=250
                )
                
                # NÚT XUẤT FILE DOCX CHUẨN MỰC
                st.markdown("<br>", unsafe_allow_html=True)
                final_df = edited_df[edited_df['Chọn'] == True].drop(columns=['Chọn'])
                
                doc_data = generate_word_report(
                    final_df, 
                    topic, 
                    author, 
                    ai.get("tu_khoa_vn", []), 
                    ai.get("tu_khoa_en", [])
                )
                
                st.download_button(
                    label="📥 Xuất báo cáo DOCX chuẩn ULAW",
                    data=doc_data,
                    file_name="Bao_cao_Tra_cuu_USRS.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    type="primary"
                )
            else:
                st.warning("Không tìm thấy kết quả phù hợp.")
